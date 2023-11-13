import streamlit as st
import openai
from docx import Document

st.set_page_config(
    page_title="Meeting Minutes Generator",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("## :gray[Meeting Minutes Generator]")

with st.form(key='api_form'):
    st.markdown("""
    Enter your OpenAI API token :red[*]
    """)
    api_key = st.text_input("Enter your OpenAI API token:", type='password', key = 'token', label_visibility='collapsed')
    st.form_submit_button("SUBMIT")

if not (api_key.startswith('sk-') and len(api_key) == 51):
    st.warning('Please enter your credentials!', icon = '⚠️')
else:
    st.success("Proceed to use the app!", icon = '✅')

transcribe, generate, minutes = st.tabs([":blue[Transcribe]", ":blue[Generate]", "Meeting Minutes"], )

@st.cache_data(show_spinner=False)
def transcribe_audio(audio_file, language):
    openai.api_key = api_key
    transcription = openai.audio.transcriptions.create(
            model = "whisper-1", 
            file = audio_file,
            language=language
        )
    return transcription.text

global r
r = ""
with transcribe:
    st.subheader("Transcribes audio into the input language")

    uploaded_audio = st.file_uploader("Choose a audio file", key='audio_input', disabled=not api_key, type=['flac', 'mp3', 'mp4', 'mpeg', 'mpga', 'm4a', 'ogg', 'wav', 'webm'])
    st.text("Select language")
    # language in ISO-639-1 format to improve the accuracy and latency of the transcription
    language_dict = {
        "English": "en",
        "Spanish": "es",
        "French": "fr",
        "German": "de",
        "Italian": "it",
        "Portuguese": "pt"
    }
    col1, col2 = st.columns(2)
    with col1:
        language = st.selectbox("Select language", ("English", "Spanish", "French", "German", "Italian", "Portuguese"), index=0, disabled=not uploaded_audio, placeholder="The language to transcribe the audio into", label_visibility='collapsed')
    with col2:
        t_button = st.button("Transcribe", type="primary", disabled=not uploaded_audio)
    st.divider()

    if uploaded_audio and t_button:
        progress_bar = st.progress(0, "Operation in progress. Please wait...")
        result = transcribe_audio(uploaded_audio, language_dict[language])
        r = result
        progress_bar.progress(100)
        st.text_area(label="", value=result, label_visibility='collapsed', key="t_1")
    if r:
        with open('transcript.txt', 'w') as f:
            f.write(r)
            st.download_button("Download transcript", data = r, mime="text/plain", file_name="transcript.txt")

@st.cache_data(show_spinner=False)
def generate_audio(txt, voice, speed):
    openai.api_key = api_key
    speech_file_path = "speech.mp3"
    response = openai.audio.speech.create(
        model="tts-1",
        voice=voice,
        input=txt,
        speed=speed
    )

    response.stream_to_file(speech_file_path)
    return speech_file_path

global a
a = ''
with generate:
    st.subheader("Generates audio from the input text")

    txt = st.text_area("Text to generate audio for", disabled=not api_key, max_chars=4096)
    if txt is not None:
        col1, col2 = st.columns(2)
        with col1:
            voice = st.selectbox("Select voice", ("alloy", "echo", "fable", "onyx", "nova", "shimmer"), index=4, disabled=not txt, help="The voice to use when generating the audio.")
        with col2:
            speed = st.slider("Specify speed", min_value=0.25, max_value=4.0, value=1.0, disabled=not txt, step=0.05, help = "The speed of the generated audio.", )
        g_button = st.button("Generate", type="primary", disabled=not txt)
    st.divider()
    if txt is not None and g_button:
        progress_bar = st.progress(0, "Operation in progress. Please wait...")
        audio_file_path = generate_audio(txt, voice, speed)
        a = audio_file_path
        progress_bar.progress(100)
        st.audio(audio_file_path, format="audio/mp3")
    if a:
        with open('speech.mp3', 'rb') as f:
            audio_content = f.read()
        st.download_button("Download audio", data = audio_content, mime="audio/mp3", file_name="speech.mp3")

# Summary extraction
@st.cache_data(show_spinner=False)
def abstract_summary_extraction(transcription):
    openai.api_key = api_key
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

# Key points extraction
@st.cache_data(show_spinner=False)
def key_points_extraction(transcription):
    openai.api_key = api_key
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content


# Action item extraction
@st.cache_data(show_spinner=False)
def action_item_extraction(transcription):
    openai.api_key = api_key
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

# Sentiment analysis
@st.cache_data(show_spinner=False)
def sentiment_analysis(transcription):
    openai.api_key = api_key
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

@st.cache_data(show_spinner=False)
def meeting_minutes(transcription, options):
    minutes = {
        'abstract_summary': None,
        'key_points': None,
        'action_items': None,
        'sentiment': None
    }
    minutes['abstract_summary'] = abstract_summary_extraction(transcription)
    if 'Key points' in options:
        minutes['key_points'] = key_points_extraction(transcription)
    if 'Action items' in options:
        minutes['action_items'] = action_item_extraction(transcription)
    if 'Sentiment' in options:
        minutes['sentiment'] = sentiment_analysis(transcription)
    
    return minutes
    
# Exporting meeting minutes
def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)

global m
m = {}
with minutes:
    st.subheader("Generate meeting minutes")
    uploaded_audio = st.file_uploader("Choose a audio file", key='audio_input_2', disabled=not api_key, type=['flac', 'mp3', 'mp4', 'mpeg', 'mpga', 'm4a', 'ogg', 'wav', 'webm'])
    st.text("Analyze")
    col1, col2 = st.columns(2)
    with col1:
        options = st.multiselect(
            'Analyze',
            ['Key points', 'Action items', 'Sentiment'],
            default=None,
            max_selections=2,
            disabled=not uploaded_audio,
            label_visibility='collapsed',
            placeholder='Optional analysis',
        )
    with col2:
        g_button = st.button("Generate", type="primary", disabled=not uploaded_audio, key='g_2')
    st.divider()
    if uploaded_audio and g_button:
        with st.spinner("Operation in progress. Please wait..."):
            # Transcribe in English
            result = transcribe_audio(uploaded_audio, 'en')
            m_result = meeting_minutes(result, options)
            m = m_result
            st.text('Abstract summary')
            st.text_area(label="", value=m_result['abstract_summary'], label_visibility='collapsed')
            if 'Key points' in options:
                st.text('Key points')
                st.text_area(label="", value=m_result['key_points'], label_visibility='collapsed')
            if 'Action items' in options:
                st.text('Action items')
                st.text_area(label="", value=m_result['action_items'], label_visibility='collapsed')
            if 'Sentiment' in options:
                st.text('Sentiment')
                st.text_area(label="", value=m_result['sentiment'], label_visibility='collapsed')

    if m:
        save_as_docx(m, 'meeting_minutes.docx')
        with open('meeting_minutes.docx', 'rb') as f:
            st.download_button("Download Meeting Minutes", data=f.read(), mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document', file_name='meeting_minutes.docx')